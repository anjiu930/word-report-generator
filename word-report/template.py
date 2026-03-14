# -*- coding: utf-8 -*-
"""
项目报告生成脚本
字数要求：6800字左右
"""
from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
import os
import copy

OUTPUT_PATH = r"D:\report\word\项目报告.docx"
TEMPLATE_PATH = r"C:\Users\31930\.claude\skills\word-report\examples\课程设计报告书模板.docx"

# 封面信息
COVER_INFO = {
    "school": "学校名称",
    "college": "学院名称",
    "course": "课程名称",
    "title": "项目标题",
    "major_class": "专业班级",
    "name": "姓名",
    "student_id": "学号",
    "teacher": "授课教师",
    "year": "年份",
    "month": "",
    "day": ""
}
# ================================================
def set_chinese_font(run, font_name='宋体', font_size=Pt(10.5), bold=False, italic=False, underline=False):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    # 设置中文字体
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)
def set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         first_line_indent=None, line_spacing=1.25,
                         space_before=Pt(0), space_after=Pt(0)):
    """设置段落格式"""
    pf = paragraph.paragraph_format
    pf.alignment = alignment
    pf.line_spacing = line_spacing
    pf.space_before = space_before
    pf.space_after = space_after
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
def enable_trailing_space_underline(doc):
    """开启'为尾部空格添加下划线'设置（备用函数，供自定义封面使用）"""
    settings = doc.settings.element
    compat = settings.find(qn('w:compat'))
    if compat is None:
        compat = OxmlElement('w:compat')
        settings.append(compat)
    ul_trail_space = OxmlElement('w:ulTrailSpace')
    compat.append(ul_trail_space)
def create_cover_page(doc, info):
    """创建封面页（备用函数，模板已自带封面时无需调用）"""
    # 学校名称
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    run = p.add_run(info["school"])
    set_chinese_font(run, font_size=Pt(22), bold=True)

    # 学院名称
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    run = p.add_run(info["college"])
    set_chinese_font(run, font_size=Pt(18), bold=True)

    # 空行
    doc.add_paragraph()
    doc.add_paragraph()

    # 课程名称
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    run = p.add_run("课程名称：")
    set_chinese_font(run, font_size=Pt(14))
    run = p.add_run("     " + info["course"])
    set_chinese_font(run, font_size=Pt(14), underline=True)

    # 空行
    doc.add_paragraph()

    # 封面表格
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表格宽度
    table.columns[0].width = Cm(3)
    table.columns[1].width = Cm(6)

    # 表格内容
    table_data = [
        ("题    目：", info["title"]),
        ("专业班级：", info["major_class"]),
        ("姓    名：", info["name"]),
        ("学    号：", info["student_id"]),
        ("授课教师：", info["teacher"]),
        ("成    绩：", "")
    ]

    for i, (label, value) in enumerate(table_data):
        # 标签单元格
        cell_label = table.cell(i, 0)
        p = cell_label.paragraphs[0]
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=1.5)
        run = p.add_run(label)
        set_chinese_font(run, font_size=Pt(14))

        # 值单元格（带下划线）
        cell_value = table.cell(i, 1)
        p = cell_value.paragraphs[0]
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5)
        # 使用空格填充以显示下划线
        value_text = value if value else "          "
        value_text = value_text.ljust(20)  # 确保有足够的空格显示下划线
        run = p.add_run(value_text)
        set_chinese_font(run, font_size=Pt(14), underline=True)

    # 空行
    doc.add_paragraph()
    doc.add_paragraph()

    # 日期
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    date_text = f"{info['year']} 年    月    日"
    run = p.add_run(date_text)
    set_chinese_font(run, font_size=Pt(14))

    # 分页符
    doc.add_page_break()
def add_chapter_title(doc, text):
    """添加一级标题（章标题）"""
    p = doc.add_paragraph(text, style='Heading 1')
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         space_before=Pt(0), space_after=Pt(4), line_spacing=1.25)
    for run in p.runs:
        set_chinese_font(run, font_size=Pt(15), bold=True)
    return p
def add_section_title(doc, text):
    """添加二级标题（小节标题）"""
    p = doc.add_paragraph(text, style='Heading 2')
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         space_before=Pt(2), space_after=Pt(0), line_spacing=1.25)
    for run in p.runs:
        set_chinese_font(run, font_size=Pt(12), bold=True)
    return p
def add_summary_paragraph(doc, text):
    """添加总领段落（无缩进）"""
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         first_line_indent=Cm(0), line_spacing=1.25,
                         space_before=Pt(0), space_after=Pt(0))
    run = p.add_run(text)
    set_chinese_font(run, font_size=Pt(10.5))
    return p
def add_body_paragraph(doc, text):
    """添加正文段落（首行缩进2字符）"""
    p = doc.add_paragraph()
    # 2字符缩进，五号字体约为0.74cm
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         first_line_indent=Cm(0.74), line_spacing=1.25,
                         space_before=Pt(0), space_after=Pt(0))
    run = p.add_run(text)
    set_chinese_font(run, font_size=Pt(10.5))
    return p
def add_code_block(doc, code_text):
    """添加代码块"""
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         first_line_indent=Cm(0), line_spacing=1.0,
                         space_before=Pt(3), space_after=Pt(3))
    run = p.add_run(code_text)
    set_chinese_font(run, font_name='宋体', font_size=Pt(9))
    return p
def add_image_placeholder(doc, description, fig_num):
    """添加图片占位符（在设计章节中使用，上一段应有引用说明）"""
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.25)
    run = p.add_run(f"图{fig_num} {description}")
    set_chinese_font(run, font_size=Pt(10.5))
    return p
def add_image_placeholder_with_ref(doc, description, fig_num, ref_text=None):
    """添加带引用的图片占位符（备用函数，可替代手动写引用段落 + add_image_placeholder 的组合）

    Args:
        doc: Document 对象
        description: 图片描述
        fig_num: 图片编号
        ref_text: 引用文字（可选，默认生成"如图X所示"）
    """
    # 如果没有提供引用文字，自动生成
    if ref_text is None:
        ref_text = f"如图{fig_num}所示"

    # 先添加引用段落
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         first_line_indent=Cm(0.74), line_spacing=1.25,
                         space_before=Pt(0), space_after=Pt(0))
    run = p.add_run(ref_text)
    set_chinese_font(run, font_size=Pt(10.5))

    # 添加图片占位符
    add_image_placeholder(doc, description, fig_num)
    return p
def add_code_with_description(doc, description, code_text, code_language="c"):
    """添加带描述的代码块（先描述功能，再展示代码）

    Args:
        doc: Document 对象
        description: 功能描述文字
        code_text: 代码文本内容
        code_language: 代码语言（用于注释）
    """
    # 先添加功能描述段落
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         first_line_indent=Cm(0.74), line_spacing=1.25,
                         space_before=Pt(0), space_after=Pt(0))
    run = p.add_run(description)
    set_chinese_font(run, font_size=Pt(10.5))

    # 添加代码块
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         first_line_indent=Cm(0), line_spacing=1.0,
                         space_before=Pt(3), space_after=Pt(3))
    run = p.add_run(code_text)
    set_chinese_font(run, font_name='宋体', font_size=Pt(9))
    return p
def add_references_title(doc):
    """添加参考文献标题"""
    # 分页符
    doc.add_page_break()

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(0), space_after=Pt(4), line_spacing=1.25)
    run = p.add_run("参考文献")
    set_chinese_font(run, font_size=Pt(15), bold=True)
    return p
def add_reference_item(doc, text):
    """添加参考文献条目"""
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         first_line_indent=Cm(0), line_spacing=1.25,
                         space_before=Pt(0), space_after=Pt(0))
    run = p.add_run(text)
    set_chinese_font(run, font_size=Pt(10.5))
    return p
def find_toc_end(doc, default=35):
    """自动查找目录结束位置（第一个章标题前一段即为目录末尾）

    Args:
        doc: Document 对象
        default: 如果未找到章标题，返回的默认索引值

    Returns:
        int: 目录结束的段落索引
    """
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style_name = para.style.name if para.style else ""
        # 通过样式或文本前缀识别第一个章标题
        if style_name.startswith('Heading 1') or style_name == '标题 1':
            if text:
                return max(i - 1, 0)
        if text.startswith("第") and "章" in text[:6]:
            return max(i - 1, 0)
    return default
def remove_content_after_toc(doc, toc_end_index=35):
    """删除目录之后的所有段落，保留所有表格

    Args:
        doc: Document 对象
        toc_end_index: 目录结束的段落索引（从0开始）
            - 默认值为35，适用于 examples/课程设计报告书模板.docx
            - 如使用自定义模板，需要先分析模板结构确定正确的索引值
            - 可通过遍历 doc.paragraphs 查找目录结束位置

    注意：
        此函数只删除段落，不删除表格。
        如需删除模板中的示例表格，请使用 remove_table_by_index() 函数。
    """
    for i in range(len(doc.paragraphs) - 1, toc_end_index, -1):
        para = doc.paragraphs[i]
        para._element.getparent().remove(para._element)
def remove_table_by_index(doc, table_index):
    """删除指定索引的表格

    Args:
        doc: Document 对象
        table_index: 要删除的表格索引（从0开始）

    使用场景：
        当模板中包含示例表格（如课程设计报告书模板中的 Altera 宏功能单元表）时，
        需要在生成报告前删除这些不需要的表格。
    """
    if len(doc.tables) > table_index:
        table = doc.tables[table_index]
        table._element.getparent().remove(table._element)
        print(f"已删除表格（索引 {table_index}）")
def move_table_to_end(doc, table_index):
    """将指定表格移到文档末尾

    Args:
        doc: Document 对象
        table_index: 要移动的表格索引（从0开始）

    使用场景：
        将评分表从模板正文位置移动到参考文献之后，作为报告的最后一页。
    """
    table = doc.tables[table_index]
    table_element = table._element

    # 从原位置移除
    table_element.getparent().remove(table_element)

    # 添加到文档末尾
    doc._element.body.append(table_element)
def create_sample_report():
    """创建项目报告 - 字数约6800字

    模板处理流程：
    1. 读取模板文档
    2. 删除目录之后的段落（保留表格）
    3. 【可选】删除模板中的示例表格
    4. 添加正文内容
    5. 添加参考文献
    6. 将评分表移动到文档末尾

    注意事项：
    - 如果模板中有示例表格，需要调用 remove_table_by_index() 删除
    - 删除表格后，剩余表格的索引会发生变化
    """
    print("正在读取模板文档...")
    doc = Document(TEMPLATE_PATH)

    # 步骤1：自动检测目录结束位置，然后删除目录之后的段落（保留表格）
    print("正在清理模板正文...")
    toc_end = find_toc_end(doc, default=35)
    print(f"  检测到目录结束位置：段落索引 {toc_end}")
    remove_content_after_toc(doc, toc_end_index=toc_end)

    # 步骤2：【可选】删除模板中的示例表格
    # 如果模板中有不需要的示例表格，取消下面的注释并调整索引
    # 示例：删除第三个表格（索引为2）
    # remove_table_by_index(doc, 2)

    # doc.add_page_break()

    # 步骤2：用代码创建正文内容
    print("正在生成正文内容...")

    # ========== 第1章 设计题目 ==========

    add_chapter_title(doc, "第1章 设计题目")
    add_body_paragraph(doc, "项目背景介绍，描述项目来源和应用场景，行业现状分析，介绍项目的宏观背景、行业痛点、技术发展趋势等。")

    add_body_paragraph(doc, "本课程设计旨在解决上述问题，通过软硬件协同的方式开发一套功能完整、性能稳定的系统解决方案，力求在功能实现、性能优化和用户体验三个维度上达到预期设计目标。")    add_body_paragraph(doc, "系统架构概述段落，描述技术栈和整体架构，概述系统的核心技术和设计思路。")

    doc.add_paragraph()

    # ========== 第2章 开发环境 ==========
    add_chapter_title(doc, "第2章 开发环境")
    add_section_title(doc, "2.1 硬件环境")
    add_body_paragraph(doc, "硬件环境描述，包括处理器、内存、硬盘、测试设备等配置信息。")

    add_section_title(doc, "2.2 软件环境")
    add_body_paragraph(doc, "软件环境描述，包括操作系统、开发工具、服务器、数据库等软件配置。")

    doc.add_paragraph()

    # ========== 第3章 开发工具 ==========
    add_chapter_title(doc, "第3章 开发工具")
    add_section_title(doc, "3.1 开发工具1")
    add_body_paragraph(doc, "开发工具1详细介绍，描述其功能定位、核心特性、版本选择理由，以及在本项目中承担的具体角色和使用方式。")

    add_section_title(doc, "3.2 开发工具2")
    add_body_paragraph(doc, "开发工具2是本项目不可或缺的辅助工具，承担了代码编译和项目构建的核心工作。")
    add_body_paragraph(doc, "开发工具2功能介绍，描述其在项目中的具体应用场景。")

    add_section_title(doc, "3.3 开发工具3")
    add_body_paragraph(doc, "开发工具3介绍，描述其在版本控制和团队协作方面提供的关键支持与核心特点。")

    add_section_title(doc, "3.4 开发工具4")
    add_body_paragraph(doc, "开发工具4在系统调试和性能分析阶段发挥了重要作用，显著提升了问题定位效率。")
    add_body_paragraph(doc, "开发工具4详细说明，介绍其工作原理、配置方法和项目中的实际使用效果。")

    add_section_title(doc, "3.5 开发工具5")
    add_body_paragraph(doc, "开发工具5介绍，描述其在项目部署和运维管理方面提供的自动化支持与便利功能。")

    doc.add_paragraph()

    # ========== 第4章 完成时间 ==========
    add_chapter_title(doc, "第4章 完成时间")
    add_body_paragraph(doc, "项目开发周期描述，包括起止时间、各阶段工作内容等。")
    doc.add_paragraph()

    # ========== 第5章 设计思想 ==========
    add_chapter_title(doc, "第5章 设计思想")
    add_section_title(doc, "5.1 架构设计")
    add_body_paragraph(doc, "整体架构详细描述，从底层到顶层逐层介绍系统的层级划分、各层职责、层间接口定义，以及为什么选择这种分层方式。")

    add_body_paragraph(doc, "这种分层架构有效降低了各模块间的耦合度，当某一层的实现细节发生变化时，只需修改该层内部代码而不会影响其他层的正常运行，显著提升了系统的可维护性。")    add_body_paragraph(doc, "具体设计模式和实现方式描述，介绍所采用的关键设计模式及其在本项目中的应用场景和好处。")

    add_section_title(doc, "5.2 设计模式或架构风格")
    add_body_paragraph(doc, "架构风格选择理由及对比分析，结合项目需求说明所选架构的优势和适用场景。")

    add_body_paragraph(doc, "该方案经过多个实际项目的开发验证，在扩展性和稳定性方面均表现出色，能够适配不同规模的业务场景，具备良好的跨平台兼容性和长期演进能力。")    add_section_title(doc, "5.3 关键技术或策略")
    add_body_paragraph(doc, "关键技术深入分析，介绍核心算法或技术的选型理由、工作原理和实现策略。")

    add_section_title(doc, "5.4 核心机制")
    add_body_paragraph(doc, "核心机制的技术实现细节，阐述其工作原理以及在系统中承担的关键职责。")

    add_body_paragraph(doc, "该机制从数据采集、信号处理和控制输出等多个维度保障了系统运行的稳定性与实时性，是整体技术方案中不可替代的核心支撑环节。")    doc.add_paragraph()

    # ========== 第6章 设计过程及设计步骤 ==========
    add_chapter_title(doc, "第6章 设计过程及设计步骤")
    add_section_title(doc, "6.1 硬件设计")
    add_body_paragraph(doc, "本系统的硬件设计基于项目需求文档，选择了适合的硬件平台和外围电路。以下对各核心硬件模块进行详细说明。")
    # 硬件1：微控制器/处理器
    add_section_title(doc, "6.1.1 主控芯片")
    add_body_paragraph(doc, "选用XXX芯片作为系统主控，该芯片具有高性能、低功耗的特点，内部集成了ADC、PWM、UART等丰富的外设资源，能够满足本项目的控制需求。如图1所示，该芯片采用XXX封装，外部引脚功能定义清晰，便于电路设计和调试。")

    add_body_paragraph(doc, "芯片主频可达XXX MHz，同时内置多种低功耗运行模式和硬件看门狗定时器，在保证高速运算能力的前提下有效控制了系统整体功耗。")    add_image_placeholder(doc, "主控芯片原理图", 1)

    # 硬件2：传感器
    add_section_title(doc, "6.1.2 传感器模块")
    add_body_paragraph(doc, "该传感器凭借出色的测量稳定性和批次一致性，在工业过程控制、环境监测和民用智能家居等多个领域积累了大量成功的应用案例。")    add_body_paragraph(doc, "选用XXX传感器用于采集环境参数，该传感器具有精度高、响应快、功耗低等优点。传感器通过I2C/SPI接口与主控芯片通信，采样频率可达XXX Hz。如图2所示，传感器模块集成了信号调理电路，可直接输出数字信号。")

    add_body_paragraph(doc, "传感器工作电压范围与主控芯片的供电规格完全匹配，二者之间无需额外增加电平转换电路，从而简化了整体硬件设计并降低了制造成本。")    add_image_placeholder(doc, "传感器模块原理图", 2)

    # 硬件3：通信模块
    add_section_title(doc, "6.1.3 通信模块")
    add_body_paragraph(doc, "选用XXX无线通信模块实现数据传输，该模块支持XXX协议，传输距离可达XXX米，传输速率满足项目要求。如图3所示，通信模块采用邮票孔封装，可直接焊接在PCB上。")

    add_body_paragraph(doc, "通信模块通过UART串口接口与主控芯片建立物理连接，上层使用标准AT指令集进行工作模式配置、网络参数设定以及双向数据收发的全流程控制。")    add_image_placeholder(doc, "通信模块原理图", 3)

    # 硬件4：执行器/显示等
    add_section_title(doc, "6.1.4 执行机构")
    add_body_paragraph(doc, "选用XXX执行器实现控制输出，该执行器响应速度快、驱动力强，能够可靠地完成控制任务。如图4所示，执行器驱动电路采用光耦隔离设计，有效保护主控芯片。执行机构由主控芯片的GPIO口通过驱动电路控制，支持PWM调速和方向控制。")

    add_image_placeholder(doc, "执行机构原理图", 4)

    # 硬件连接总图
    add_section_title(doc, "6.1.5 系统硬件连接图")
    add_body_paragraph(doc, "各硬件模块通过统一的电源管理系统进行集中供电和分路保护，整体采用模块化设计理念，各功能单元之间相互独立，便于后期维护和功能升级。")    add_body_paragraph(doc, "如图5所示，系统采用总线式连接架构，各模块之间通过排针排母连接，便于扩展和更换。整体硬件布局兼顾信号完整性与散热需求。")

    add_image_placeholder(doc, "系统硬件连接图", 5)

    add_section_title(doc, "6.2 数据库设计")
    add_body_paragraph(doc, "数据库设计阶段介绍，包括概念设计、逻辑设计、物理设计。数据库设计是系统开发的重要环节，良好的数据库设计能够提高数据存取效率，保证数据完整性。如图6所示，本系统采用ER图进行概念模型设计。")

    add_body_paragraph(doc, "具体表结构的设计紧密围绕实际业务需求展开，合理规划了主要数据表的名称、核心字段、数据类型以及表间的外键约束和索引策略。")    add_image_placeholder(doc, "数据库ER图", 6)
    add_section_title(doc, "6.3 功能模块开发")
    add_body_paragraph(doc, "系统整体采用模块化设计思想，按照高内聚低耦合的原则将复杂功能拆分为若干独立子模块，各模块之间通过标准化接口进行数据交互。")    add_body_paragraph(doc, "功能模块开发概述。系统功能模块划分如图7所示，将系统划分为多个相对独立的功能模块，各模块之间通过定义良好的接口进行通信。")

    # 软件实现 - 传感器数据采集
    add_section_title(doc, "6.3.1 传感器数据采集实现")
    add_code_with_description(doc,
        "传感器数据采集通过I2C总线读取，首先初始化I2C配置，然后发送设备地址和寄存器地址，最后读取数据。",
        "// 传感器数据读取函数\nuint16_t Sensor_ReadData(uint8_t reg_addr)\n{\n    uint8_t data[2];\n    I2C_Start();                    // 发送起始信号\n    I2C_SendByte(SENSOR_ADDR << 1); // 发送从机地址（写）\n    I2C_SendByte(reg_addr);         // 发送寄存器地址\n    I2C_Start();                    // 重新发送起始信号\n    I2C_SendByte((SENSOR_ADDR << 1) | 0x01); // 发送从机地址（读）\n    data[0] = I2C_ReceiveByte(ACK); // 读取高字节\n    data[1] = I2C_ReceiveByte(NACK);// 读取低字节\n    I2C_Stop();\n    return (data[0] << 8) | data[1];\n}")

    # 软件实现 - 通信模块
    add_section_title(doc, "6.3.2 无线通信实现")
    add_code_with_description(doc,
        "无线通信模块采用串口通信方式，通过AT指令控制模块进行数据传输。以下为初始化和发送数据的核心代码。",
        "// 无线通信初始化\nvoid Wireless_Init(void)\n{\n    USART_Init(9600);               // 初始化串口波特率9600\n    delay_ms(100);\n    USART_SendString(\"AT+RST\\r\\n\"); // 复位模块\n    delay_ms(1000);\n    USART_SendString(\"AT+CWMODE=1\\r\\n\"); // 设置为Station模式\n    delay_ms(500);\n}\n\n// 数据发送函数\nuint8_t Wireless_SendData(uint8_t *data, uint16_t len)\n{\n    char cmd[64];\n    sprintf(cmd, \"AT+CIPSEND=%d\\r\\n\", len);\n    USART_SendString(cmd);\n    delay_ms(100);\n    USART_SendData(data, len);\n    return 0;\n}")

    # 软件实现 - 控制算法
    add_section_title(doc, "6.3.3 控制算法实现")
    add_code_with_description(doc,
        "系统采用PID控制算法实现精确控制，通过比例、积分、微分三个环节的加权组合，实现良好的动态响应和稳态性能。",
        "// PID控制算法\nfloat PID_Calculate(PID_TypeDef *pid, float setpoint, float feedback)\n{\n    float error = setpoint - feedback;  // 计算误差\n    \n    // 比例环节\n    pid->kp_output = pid->kp * error;\n    \n    // 积分环节（带积分限幅）\n    pid->integral += error * pid->ki;\n    if(pid->integral > pid->integral_max) pid->integral = pid->integral_max;\n    if(pid->integral < pid->integral_min) pid->integral = pid->integral_min;\n    \n    // 微分环节（低通滤波）\n    pid->derivative = pid->kd * (error - pid->prev_error);\n    pid->prev_error = error;\n    \n    // 总输出\n    pid->output = pid->kp_output + pid->integral + pid->derivative;\n    \n    // 输出限幅\n    if(pid->output > pid->output_max) pid->output = pid->output_max;\n    if(pid->output < pid->output_min) pid->output = pid->output_min;\n    \n    return pid->output;\n}")

    add_body_paragraph(doc, "前端界面开发详细描述，包括页面布局设计、交互逻辑实现以及用户体验优化等方面。")

    add_body_paragraph(doc, "该模块对外暴露的接口采用了简洁清晰的设计风格，通过统一的数据格式规范和参数校验机制，有效降低了上层业务调用时的集成复杂度。")    add_body_paragraph(doc, "后端服务开发详细描述，涵盖业务逻辑处理、数据持久化操作及接口安全防护等内容。")

    add_body_paragraph(doc, "该功能模块的实现充分体现了面向对象的设计理念，通过合理运用封装、继承和多态等机制显著提升了代码的复用率和可测试性。")    add_image_placeholder(doc, "系统功能模块图", 7)
    add_section_title(doc, "6.4 接口联调")
    add_body_paragraph(doc, "前后端联调工作是系统集成过程中的关键环节，接口协议的一致性和数据格式的正确性直接关系到各模块能否在真实环境下协同运行。")    add_body_paragraph(doc, "接口联调过程描述，包括测试方法、遇到的问题和解决方案。各模块接口按照预定义的协议进行对接，通过逐步集成的方式完成整体系统联调。")

    doc.add_paragraph()

    # ========== 第7章 测试运行 ==========
    # 测试章节
    add_chapter_title(doc, "第7章 测试运行")
    add_section_title(doc, "7.1 功能测试")
    add_body_paragraph(doc, "功能测试范围和方法描述。测试内容包括各功能模块的输入输出验证、边界条件测试、异常处理测试等。")

    add_body_paragraph(doc, "经过多轮反复测试验证，各项功能测试结果均达到了设计文档中规定的预期指标，系统在正常工况和边界条件下均能保持稳定运行。")    add_body_paragraph(doc, "具体测试过程和结果描述。如图8所示为系统登录功能测试界面，验证了正常登录、错误密码登录、SQL注入防护等功能。")

    add_image_placeholder(doc, "功能测试界面", 8)
    add_body_paragraph(doc, "数据查询模块的核心功能同样通过了全面而系统的验证，包括模糊匹配、精确检索和条件过滤在内的各项指标均符合预期要求。")    add_body_paragraph(doc, "如图9所示为数据查询功能测试结果，验证了模糊查询、分页显示、排序功能等核心功能的正确性。")

    add_image_placeholder(doc, "查询功能测试结果", 9)
    add_section_title(doc, "7.2 性能测试")
    add_body_paragraph(doc, "性能测试指标和结果描述。主要测试了系统的响应时间、并发处理能力、资源占用等指标。如图10所示为压力测试结果。")

    add_image_placeholder(doc, "性能测试结果", 10)
    add_section_title(doc, "7.3 运行结果")
    add_body_paragraph(doc, "部署完成后经过连续运行观察，系统整体表现稳定可靠，各功能模块在不同负载条件下均能保持快速响应且未出现异常中断。")    add_body_paragraph(doc, "系统运行情况和效果描述。如图11所示为系统实际运行界面，各功能模块运行正常，满足设计要求。")

    add_image_placeholder(doc, "系统运行界面", 11)
    doc.add_paragraph()

    # ========== 第8章 评价与修订 ==========
    # 评价章节
    add_chapter_title(doc, "第8章 评价与修订")
    add_section_title(doc, "8.1 功能评价")
    add_body_paragraph(doc, "功能完成情况评价，对照需求文档逐项检查各功能模块的实现程度和运行效果。")

    add_section_title(doc, "8.2 技术评价")
    add_body_paragraph(doc, "技术实现评价，从代码质量、架构合理性、性能表现等维度对项目进行综合评估。")

    add_section_title(doc, "8.3 存在问题")
    add_body_paragraph(doc, "系统在部分应用场景下仍然存在一些不足之处，例如在高并发请求环境下响应速度有所下降，部分边界条件的异常处理机制也不够完善，这些问题需要在后续版本的迭代开发中作为重点方向持续关注和改进。")    add_body_paragraph(doc, "具体问题分析和原因探讨，从技术选型、开发流程和测试覆盖等角度深入剖析。")

    add_section_title(doc, "8.4 改进建议")
    add_body_paragraph(doc, "后续改进方向和优化建议，结合当前技术发展趋势提出切实可行的迭代方案。")

    doc.add_paragraph()

    # ========== 第9章 设计体会 ==========
    # 体会章节
    add_chapter_title(doc, "第9章 设计体会")
    add_body_paragraph(doc, "技术和能力提升收获，回顾整个项目开发过程中的技术成长和综合能力提升。")

    add_body_paragraph(doc, "这次课程设计让我在理论与实践的结合上收获颇丰，通过亲手搭建完整的系统架构并逐步调试各个功能模块，对专业知识有了更深层次的理解，也切实感受到了将书本概念转化为可运行程序的成就感。")    add_body_paragraph(doc, "核心技术学习体会和问题解决经历，记录在开发过程中遇到的典型难题及攻克方法。")

    add_body_paragraph(doc, "通过这次实践深刻认识到，工程实践的价值远超单纯的课堂理论学习所能带来的积累，只有在真实项目中反复调试和优化，才能真正掌握从需求分析到系统部署的完整开发流程，这是任何教科书都无法替代的宝贵经验。")    add_body_paragraph(doc, "不足之处和未来展望，反思项目中的遗憾和对未来发展方向的思考。")

    doc.add_paragraph()

    # ========== 参考文献 ==========
    add_references_title(doc)
    add_reference_item(doc, "[1] 作者. 书名[M]. 出版社: 年份.")
    add_reference_item(doc, "[2] 作者. 书名[M]. 出版社: 年份.")
    add_reference_item(doc, "[3] 作者. 书名[M]. 出版社: 年份.")
    add_reference_item(doc, "[4] 作者. 书名[M]. 出版社: 年份.")
    add_reference_item(doc, "[5] 作者. 书名[M]. 出版社: 年份.")
    add_reference_item(doc, "[6] 作者. 文档名[EB/OL]. URL, 年份.")
    add_reference_item(doc, "[7] 作者. 文档名[EB/OL]. URL, 年份.")
    add_reference_item(doc, "[8] 作者. 文档名[EB/OL]. URL, 年份.")

    # 添加分页符，确保评分表在新页开始
    doc.add_page_break()

    # 将评分表移到文档末尾
    # 注意：
    # 1. 表格索引从0开始
    # 2. 如果删除了示例表格，索引会发生变化
    # 3. 通常：表格0=封面学号表，表格1=评分表
    print("正在移动评分表到文档末尾...")
    if len(doc.tables) >= 2:
        move_table_to_end(doc, 1)  # 移动第二个表格（评分表）到末尾

    # 保存文档
    print("正在保存文档...")
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"✓ 报告已生成：{OUTPUT_PATH}")
def main():
    """主程序入口"""
    print("正在生成项目报告...")
    create_sample_report()
    print("完成！")
if __name__ == "__main__":
    main()
